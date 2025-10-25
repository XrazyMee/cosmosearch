#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#
import json
from datetime import datetime
from typing import List, Dict, Any, Optional
from io import BytesIO

from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from flask import Response

from api.db.db_models import PaperSearchRecord, PaperSurveyRecord, PaperSurveyDownloadRecord, Task
from api.db.services.common_service import CommonService
from api.db.services.llm_service import LLMBundle
from api.db.services.knowledgebase_service import KnowledgebaseService
from api.db.services.document_service import DocumentService
from api.db.db_utils import bulk_insert_into_db
from api import settings
from api.utils import get_uuid
from rag.nlp import search
from rag.utils.redis_conn import REDIS_CONN
from rag.settings import get_svr_queue_name

from api.db import LLMType, TenantPermission
from loguru import logger


class PaperSearchService(CommonService):
    model = PaperSearchRecord

    @classmethod
    def search_papers(
        cls,
        search_record_id: str,
        query: str,
        tenant_id: str,  # Add tenant_id parameter
        keywords_num: int = 5,
        query_num: int = 5,
        use_fuzzy: bool = True,
    ) -> Dict[str, Any]:
        """执行论文检索"""
        # 1. 提取关键词 (独立的LLM调用，不依赖RAG)
        keywords_response = cls.extract_keywords(query, tenant_id, keywords_num, query_num)

        # 2. 更新检索记录
        search_record = {
            "id": search_record_id,
            "keywords": json.dumps(keywords_response, ensure_ascii=False),
        }
        cls.model.update(search_record).where(cls.model.id == search_record_id).execute()

        # 3. 构建检索查询 - 合并所有关键词
        all_keywords = []
        all_keywords.extend(keywords_response.get("keyword_en", []))
        all_keywords.extend(keywords_response.get("keyword_cn", []))
        all_keywords.extend(keywords_response.get("searchquery_en", []))
        all_keywords.extend(keywords_response.get("searchquery_cn", []))

        search_query = " ".join(all_keywords)

        # 4. 执行RAGFlow检索 - 获取用户所有知识库以及所有公开知识库
        #    如果没有知识库，则返回空结果但关键词提取成功
        user_kbs = KnowledgebaseService.query(tenant_id=tenant_id)
        public_kbs = KnowledgebaseService.query(permission=TenantPermission.PUBLIC.value)

        # 合并用户知识库和公开知识库
        all_kbs = user_kbs + public_kbs
        kb_ids = [kb.id for kb in all_kbs]

        papers = []
        if kb_ids:
            # 如果有可用知识库，则执行检索
            try:
                # 获取知识库信息 - use the first kb's tenant_id to get embedding model
                first_kb = all_kbs[0] if all_kbs else None
                if not first_kb:
                    # 没有有效的知识库
                    pass
                else:
                    # 获取嵌入模型
                    embd_mdl = LLMBundle(first_kb.tenant_id, LLMType.EMBEDDING, llm_name=first_kb.embd_id)

                    # 执行检索使用RAGFlow原生检索能力
                    retriever = search.Dealer(settings.docStoreConn)

                    # 检索参数
                    page = 1
                    page_size = 30
                    similarity_threshold = 0.2
                    vector_similarity_weight = 0.3
                    top_k = 1024

                    # 执行RAGFlow检索
                    ranks = retriever.retrieval(
                        search_query,
                        embd_mdl,
                        first_kb.tenant_id,
                        kb_ids,
                        page=page,
                        page_size=page_size,
                        similarity_threshold=similarity_threshold,
                        vector_similarity_weight=vector_similarity_weight,
                        top=top_k,
                        doc_ids=None,
                        aggs=True,
                    )

                    # 5. 处理检索结果 - 将文档块信息转换为论文格式
                    processed_docs = set()

                    for chunk in ranks.get("chunks", []):
                        doc_id = chunk.get("doc_id", "")

                        # 避免重复处理同一文档
                        if doc_id in processed_docs:
                            continue
                        processed_docs.add(doc_id)

                        # 从数据库获取文档详细信息
                        doc_info = DocumentService.get_by_id(doc_id)
                        if doc_info and doc_info[0] and doc_info[1]:
                            doc = doc_info[1]

                            paper = {
                                "uid": doc.id,
                                "title": doc.name,  # 使用文档名称作为标题
                                "abstract": chunk.get("content_with_weight", "")[:500],  # 使用内容前500字符作为摘要
                                "source": "RAGFlow知识库",
                                "selected": True,
                                "similarity": chunk.get("similarity", 0.0),
                                "doc_id": doc_id,
                                "kb_id": chunk.get("kb_id", ""),
                            }

                            papers.append(paper)
            except Exception as e:
                # 如果RAG检索出错，仍然返回关键词但papers为空
                logger.error(f"RAG检索失败: {str(e)}")
                papers = []
        else:
            # 没有可用知识库，返回空结果
            logger.warning("未找到可用的知识库，返回空检索结果")
            papers = []

        # 6. 更新检索记录
        search_record_update = {"search_results": json.dumps(papers, ensure_ascii=False), "result_count": len(papers)}

        cls.model.update(search_record_update).where(cls.model.id == search_record_id).execute()

        # 7. 返回结果
        return {"search_record_id": search_record_id, "papers": papers, "keywords": keywords_response}

    @classmethod
    def _get_full_document_content(cls, doc_id: str, tenant_id: str, kb_ids: list) -> str:
        """获取文档的完整内容"""
        try:
            from rag.nlp.search import Dealer
            from api import settings
            
            # 使用检索器获取文档的所有块
            retriever = Dealer(settings.docStoreConn)
            all_chunks = retriever.chunk_list(doc_id, tenant_id, kb_ids)
            
            # 按位置排序并组合完整的文档内容
            sorted_chunks = sorted(all_chunks, key=lambda x: x.get("position_int", 0))
            
            # 组合所有块的内容
            full_content_parts = []
            for chunk in sorted_chunks:
                content = chunk.get("content_with_weight", "")
                if content:
                    full_content_parts.append(content)
            
            return "\n".join(full_content_parts)
        except Exception as e:
            logger.error(f"获取文档完整内容失败: {str(e)}")
            return ""

    @classmethod
    def search_papers_with_keywords(
        cls,
        search_record_id: str,
        query: str,
        search_query: str,  # 使用用户确认的关键词构建的搜索查询
        tenant_id: str,
        use_fuzzy: bool = True,
    ) -> Dict[str, Any]:
        """使用已确认的关键词执行论文检索"""

        # 1. 更新检索记录 - 先保存查询和关键词
        search_record = {
            "id": search_record_id,
            "keywords": json.dumps({"user_confirmed_query": search_query}, ensure_ascii=False),
        }
        cls.model.update(search_record).where(cls.model.id == search_record_id).execute()

        # 2. 执行RAGFlow检索 - 获取用户所有知识库以及所有公开知识库
        #    如果没有知识库，则返回空结果但关键词提取成功
        user_kbs = KnowledgebaseService.query(tenant_id=tenant_id)
        public_kbs = KnowledgebaseService.query(permission=TenantPermission.PUBLIC.value)

        # 确保我们处理的是知识库对象而不是布尔值
        user_kbs = [kb for kb in (user_kbs.dicts() if hasattr(user_kbs, "dicts") else user_kbs) if hasattr(kb, "id")]
        public_kbs = [kb for kb in (public_kbs.dicts() if hasattr(public_kbs, "dicts") else public_kbs) if hasattr(kb, "id")]

        # 合并用户知识库和公开知识库
        all_kbs = user_kbs + public_kbs
        kb_ids = [kb.id for kb in all_kbs]

        papers = []
        if kb_ids:
            # 如果有可用知识库，则执行检索
            try:
                # 获取知识库信息 - use the first kb's tenant_id to get embedding model
                first_kb = all_kbs[0] if all_kbs else None
                if not first_kb:
                    # 没有有效的知识库
                    pass
                else:
                    # 获取嵌入模型
                    embd_mdl = LLMBundle(first_kb.tenant_id, LLMType.EMBEDDING, llm_name=first_kb.embd_id)

                    # 执行检索使用RAGFlow原生检索能力
                    retriever = search.Dealer(settings.docStoreConn)

                    # 检索参数
                    page = 1
                    page_size = 30
                    similarity_threshold = 0.2
                    vector_similarity_weight = 0.3
                    top_k = 1024

                    # 执行RAGFlow检索
                    ranks = retriever.retrieval(
                        search_query,
                        embd_mdl,
                        first_kb.tenant_id,
                        kb_ids,
                        page=page,
                        page_size=page_size,
                        similarity_threshold=similarity_threshold,
                        vector_similarity_weight=vector_similarity_weight,
                        top=top_k,
                        doc_ids=None,
                        aggs=True,
                    )

                    # 5. 处理检索结果 - 将文档块信息转换为论文格式
                    processed_docs = set()

                    for chunk in ranks.get("chunks", []):
                        doc_id = chunk.get("doc_id", "")

                        # 避免重复处理同一文档
                        if doc_id in processed_docs:
                            continue
                        processed_docs.add(doc_id)

                        # 从数据库获取文档详细信息
                        doc_info = DocumentService.get_by_id(doc_id)
                        if doc_info and doc_info[0] and doc_info[1]:
                            doc = doc_info[1]

                            paper = {
                                "uid": doc.id,
                                "title": doc.name,  # 使用文档名称作为标题
                                "abstract": chunk.get("content_with_weight", "")[:500],  # 使用内容前500字符作为摘要
                                "source": "RAGFlow知识库",
                                "selected": True,
                                "similarity": chunk.get("similarity", 0.0),
                                "doc_id": doc_id,
                                "kb_id": chunk.get("kb_id", ""),
                            }

                            papers.append(paper)
            except Exception as e:
                # 如果RAG检索出错，仍然返回关键词但papers为空
                logger.error(f"RAG检索失败: {str(e)}")
                papers = []
        else:
            # 没有可用知识库，返回空结果
            logger.warning("未找到可用的知识库，返回空检索结果")
            papers = []

        # 6. 更新检索记录
        search_record_update = {"search_results": json.dumps(papers, ensure_ascii=False), "result_count": len(papers)}

        cls.model.update(search_record_update).where(cls.model.id == search_record_id).execute()

        # 7. 返回结果
        return {
            "search_record_id": search_record_id,
            "papers": papers,
            # 返回原始查询和用户选择的关键词
            "query": query,
            "keywords": {"user_confirmed_query": search_query},
        }

    @classmethod
    def generate_paper_summary(cls, paper: Dict[str, Any], tenant_id: str) -> str:
        """生成单篇文献简报"""
        try:
            # 获取聊天模型
            from api.db.services.llm_service import LLMBundle
            from api.db import LLMType

            chat_mdl = LLMBundle(tenant_id, LLMType.CHAT)

            # 构建简报生成提示词
            title = paper.get("title", "未知标题")
            content = paper.get("full_content", "") or paper.get("abstract", "")
            
            summary_prompt = f"""
请对以下学术论文生成一份200字左右的简报：

论文标题: {title}

论文内容:
{content[:4000]}  # 限制内容长度避免超限

请按照以下结构生成简报，确保内容准确反映原文：
1. **研究主题**：概括论文的核心研究主题
2. **主要方法**：描述论文采用的主要方法或技术
3. **关键结果**：总结论文的主要发现或实验结果
4. **创新点**：指出论文的重要创新或贡献
5. **应用价值**：说明论文的理论或实践意义
6. **局限性**：指出论文存在的局限性或待改进之处

请确保简报内容准确、详细、客观，突出论文的核心内容和贡献。
"""

            # 准备消息历史
            history = [
                {
                    "role": "user",
                    "content": summary_prompt,
                }
            ]

            # 生成简报
            chat_params = {"temperature": 0.5, "max_tokens": 1024}

            # 设置extra_body参数以支持DashScope
            extra_body = {"enable_thinking": False}
            chat_params["extra_body"] = extra_body

            summary_content = chat_mdl.chat("", history, chat_params)
            
            return summary_content
        except Exception as e:
            logger.error(f"生成单篇文献简报失败: {str(e)}")
            # 返回基于摘要的简报作为备用方案
            title = paper.get("title", "未知标题")
            abstract = paper.get("abstract", "")
            return f"标题: {title}\n摘要: {abstract}"

    @classmethod
    def generate_survey(cls, survey_record: Dict[str, Any], papers: List[Dict[str, Any]]) -> Dict[str, Any]:
        """生成论文综述"""
        # 保存初始记录
        survey_obj = PaperSurveyRecord.create(**survey_record)

        try:
            # 获取租户信息
            tenant_id = survey_record["tenant_id"]

            # 1. 如果论文没有full_content，获取完整内容
            updated_papers = []
            for paper in papers:
                updated_paper = paper.copy()
                if not updated_paper.get("full_content"):
                    # 获取文档完整内容
                    full_content = cls._get_full_document_content(updated_paper.get("doc_id", ""), tenant_id, [updated_paper.get("kb_id", "")])
                    if full_content:
                        updated_paper["full_content"] = full_content
                updated_papers.append(updated_paper)

            # 2. 生成每篇文献的简报
            paper_summaries = []
            for i, paper in enumerate(updated_papers):
                logger.info(f"正在生成第 {i+1}/{len(updated_papers)} 篇文献的简报")
                summary = cls.generate_paper_summary(paper, tenant_id)
                paper_summaries.append({
                    "title": paper.get("title", "未知标题"),
                    "summary": summary
                })

            # 3. 构建综述生成提示词（使用简报而非原文档）
            survey_prompt = cls._build_survey_prompt_from_summaries(paper_summaries)

            # 获取聊天模型
            from api.db.services.llm_service import LLMBundle
            from api.db import LLMType

            chat_mdl = LLMBundle(tenant_id, LLMType.CHAT)

            # 准备消息历史
            history = [
                {
                    "role": "user",
                    "content": survey_prompt,
                }
            ]

            # 生成综述
            # 添加兼容DashScope等API的参数
            chat_params = {"temperature": 0.7, "max_tokens": 2048}

            # 设置extra_body参数以支持DashScope
            extra_body = {"enable_thinking": False}
            chat_params["extra_body"] = extra_body

            survey_content = chat_mdl.chat(
                "",  # 系统提示词已经包含在用户输入中
                history,
                chat_params,
            )

            # 更新综述记录
            PaperSurveyRecord.update(survey_content=survey_content, status="completed").where(PaperSurveyRecord.id == survey_record["id"]).execute()

            return {"survey_id": survey_obj.id, "survey_content": survey_content, "status": "completed"}
        except Exception as e:
            # 更新综述记录状态为失败
            PaperSurveyRecord.update(status="failed").where(PaperSurveyRecord.id == survey_record["id"]).execute()

            logger.error(f"生成综述失败: {str(e)}")
            raise e

    @classmethod
    def queue_survey_task(cls, survey_record: Dict[str, Any], papers: List[Dict[str, Any]], priority: int = 0) -> str:
        """将综述生成任务加入队列

        Args:
            survey_record: 综述记录字典
            papers: 论文列表
            priority: 任务优先级，默认为0

        Returns:
            str: 任务ID
        """
        # 1. 创建 Task 记录
        task = {
            "id": get_uuid(),
            "doc_id": survey_record["id"],  # 使用 survey_id 作为 doc_id
            "from_page": 0,
            "to_page": len(papers),  # 用论文数量表示工作量
            "task_type": "paper_survey",
            "priority": priority,
            "progress": 0.0,
            "progress_msg": datetime.now().strftime("%H:%M:%S") + " 任务已创建，等待处理",
            "begin_at": datetime.now(),
        }

        # 2. 保存任务到数据库
        bulk_insert_into_db(Task, [task], True)

        # 3. 更新 survey_record 关联 task_id 和初始状态
        PaperSurveyRecord.update(task_id=task["id"], status="pending", progress=0.0, progress_msg="等待处理", process_begin_at=datetime.now()).where(
            PaperSurveyRecord.id == survey_record["id"]
        ).execute()

        # 4. 构造任务消息（包含完整数据）
        task_message = {
            **task,
            "survey_record": survey_record,
            "papers": papers,  # 传递论文数据
            "tenant_id": survey_record["tenant_id"],
        }

        # 5. 推送到 Redis 队列
        success = REDIS_CONN.queue_product(get_svr_queue_name(priority), message=task_message)

        if not success:
            logger.error(f"综述任务入队失败: task_id={task['id']}, survey_id={survey_record['id']}")
            raise Exception("无法访问 Redis，请检查 Redis 状态")

        logger.info(f"综述任务已入队: task_id={task['id']}, survey_id={survey_record['id']}, papers_count={len(papers)}")
        return task["id"]

    @classmethod
    def get_survey(cls, survey_id: str) -> Optional[Dict[str, Any]]:
        """获取综述详情"""
        try:
            survey = PaperSurveyRecord.get(PaperSurveyRecord.id == survey_id)
            return {
                "id": survey.id,
                "tenant_id": survey.tenant_id,
                "user_id": survey.user_id,
                "search_record_id": survey.search_record_id,
                "survey_content": survey.survey_content,
                "survey_title": survey.survey_title,
                "status": survey.status,
                "created_at": survey.created_at,
            }
        except Exception:
            return None

    @classmethod
    def generate_survey_doc(cls, download_record: Dict[str, Any]) -> Response:
        """生成综述文档"""
        import json
        from urllib.parse import quote

        # 保存下载记录
        PaperSurveyDownloadRecord.create(**download_record)

        # 获取综述内容
        survey = PaperSurveyRecord.get(PaperSurveyRecord.id == download_record["survey_record_id"])

        # 获取文献列表 - 优先从 survey_papers 字段获取
        papers = []
        try:
            if survey.survey_papers:
                # 优先使用综述记录中保存的文献列表
                papers = json.loads(survey.survey_papers)
                logger.info(f"从survey_papers字段获取文献列表,共{len(papers)}篇文献")
            else:
                # 兼容旧数据,尝试从搜索记录获取
                search_record = PaperSearchRecord.get_or_none(PaperSearchRecord.id == survey.search_record_id)
                if search_record and search_record.search_results:
                    papers = json.loads(search_record.search_results)
                    logger.info(f"从search_record获取文献列表,共{len(papers)}篇文献")
                else:
                    logger.warning(f"搜索记录不存在或无结果: search_record_id={survey.search_record_id}")
        except Exception as e:
            logger.error(f"获取文献列表失败: {e}")

        # 生成Word文档
        logger.info("开始生成Word文档...")
        doc_bytes = cls._generate_word_document(survey.survey_content, survey.survey_title, papers)
        logger.info(f"Word文档生成完成,大小: {len(doc_bytes)} bytes")

        # 返回文档下载响应
        from datetime import datetime

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # 使用URL编码处理中文文件名
        filename = f"{survey.survey_title}_{timestamp}.docx"
        encoded_filename = quote(filename.encode("utf-8"))

        return Response(
            doc_bytes,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}",
                "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            },
        )

    @classmethod
    def get_search_history(cls, user_id: str, tenant_id: str, page: int = 1, page_size: int = 10, keyword: str = "") -> Dict[str, Any]:
        """获取检索历史"""
        query = PaperSearchRecord.select().where((PaperSearchRecord.user_id == user_id) & (PaperSearchRecord.tenant_id == tenant_id))

        if keyword:
            query = query.where(PaperSearchRecord.query.contains(keyword))

        total = query.count()
        records = query.order_by(PaperSearchRecord.created_at.desc()).paginate(page, page_size)

        history = []
        # 为每个搜索记录添加综述信息
        for record in records:
            # 查找与该搜索记录关联的综述记录
            survey_record = PaperSurveyRecord.get_or_none(PaperSurveyRecord.search_record_id == record.id)

            record_info = {"id": record.id, "query": record.query, "result_count": record.result_count, "created_at": record.created_at}

            # 如果存在关联的综述记录，则添加综述状态信息
            if survey_record:
                record_info["survey_status"] = survey_record.status
                record_info["survey_id"] = survey_record.id

            history.append(record_info)

        return {"total": total, "records": history}

    @classmethod
    def get_survey_history(cls, user_id: str, tenant_id: str, page: int = 1, page_size: int = 10, keyword: str = "") -> Dict[str, Any]:
        """获取综述历史"""
        query = PaperSurveyRecord.select().where((PaperSurveyRecord.user_id == user_id) & (PaperSurveyRecord.tenant_id == tenant_id))

        if keyword:
            query = query.where(PaperSurveyRecord.survey_title.contains(keyword))

        total = query.count()
        records = query.order_by(PaperSurveyRecord.created_at.desc()).paginate(page, page_size)

        history = []
        for record in records:
            history.append({"id": record.id, "search_record_id": record.search_record_id, "survey_title": record.survey_title, "status": record.status, "created_at": record.created_at})

        return {"total": total, "records": history}

    @classmethod
    def delete_survey_record(cls, survey_id: str, user_id: str, tenant_id: str) -> bool:
        """删除综述记录"""
        try:
            # 验证权限
            survey = PaperSurveyRecord.get((PaperSurveyRecord.id == survey_id) & (PaperSurveyRecord.user_id == user_id) & (PaperSurveyRecord.tenant_id == tenant_id))

            # 删除记录
            survey.delete_instance()
            return True
        except Exception:
            return False

    @classmethod
    def extract_keywords(
        cls,
        question: str,
        tenant_id: str,  # Add tenant_id parameter
        keywords_num: int = 5,
        query_num: int = 5,
    ) -> Dict[str, List[str]]:
        """提取关键词"""
        # 构建提示词
        keywords_num = keywords_num if 12 >= keywords_num > 1 else 5
        query_num = query_num if query_num else keywords_num

        prompt = f"""
请根据用户的输入推断用户感兴趣的研究方向，并输出{keywords_num}个最贴近用户需求的关键词。你需要将这些关键词以JSON格式直接输出，不需要任何额外说明，格式为
  {{"keyword_en": [ 英文关键字 ], "keyword_cn": [对应的中文关键字], "searchquery_en" : [ 英文搜索句 ], "searchquery_cn" : [中文搜索句], "time_range": ["2025", "2024", "2023"] }}。
注意以下几点：
1. 关键词需要准确反映用户输入的主题。
2. 确保输出结果严格遵循指定的JSON格式。
3. 输出中不得包含任何XML标签。
4. 时间范围识别：根据用户描述的时间要求，输出年份数组，默认为空数组[]。

完成任务的步骤如下：
- 首先，仔细阅读和分析用户输入的内容，识别出核心主题或研究方向，关键词必须在该领域内。
- 其次，根据识别出的主题，提取出{keywords_num}个紧密相关的关键词，确保同时提供英文和中文版本。
- 然后，根据关键词和用户旨意提出 {query_num}个可以用于检索论文的搜索句子，确保同时提供英文和中文版本
- 识别时间范围：
  * 如果用户提到"近三年"、"最近三年"，输出["2025", "2024", "2023"]
  * 如果用户提到"近五年"、"最近五年"，输出["2025", "2024", "2023", "2022", "2021"]
  * 如果用户提到"2020年以后"、"2020年后"，输出["2025", "2024", "2023", "2022", "2021"]
  * 如果用户提到具体年份如"2023年"，输出["2023"]
  * 如果用户提到年份范围如"2022-2024"，输出["2022", "2023", "2024"]
  * 如果用户没有提到时间要求，输出[]
- 最后，按照指定的JSON格式组织输出内容，确保格式正确无误。

如果用户输入的内容不够明确，请基于常见的研究领域进行合理推测。

示例：
用户输入：我对人工智能在医疗领域的应用很感兴趣，特别是如何利用深度学习技术进行疾病诊断，希望了解近三年的研究进展。
输出：{{"keyword_en": ["Artificial Intelligence","Medical Applications","Deep Learning","Disease Diagnosis","Healthcare"], "keyword_cn": ["人工智能","医疗应用","深度学习","疾病诊断","医疗保健"] , "searchquery_en": ["Deep Learning Applications in Disease Diagnosis", "Innovations of Artificial Intelligence in Healthcare", "Artificial Intelligence-based Disease Diagnosis Technologies", "Challenges and Opportunities of Deep Learning in Healthcare"], "searchquery_cn" : [ "深度学习在疾病诊断中的应用", "人工智能在医疗保健中的创新", "基于人工智能的疾病诊断技术" ], "time_range": ["2025", "2024", "2023"]}}
"""

        # 构建system提示词和history消息
        system_prompt = prompt
        history = [
            {
                "role": "user",
                "content": question,
            }
        ]

        # 获取AI模型进行关键词提取
        try:
            from api.db.services.llm_service import LLMBundle
            from api.db import LLMType

            # 使用传入的租户ID来初始化LLMBundle
            chat_mdl = LLMBundle(tenant_id, LLMType.CHAT)

            # 调用模型进行关键词提取
            # 使用适当的参数调用chat方法
            # 添加兼容DashScope等API的参数
            chat_params = {"temperature": 0.6, "max_tokens": 512}

            # 设置extra_body参数以支持DashScope
            extra_body = {"enable_thinking": False}
            chat_params["extra_body"] = extra_body

            response = chat_mdl.chat(system_prompt, history, chat_params)

            # 解析响应
            result = cls._parse_keyword_response(response)
            return result

        except Exception as e:
            logger.error(f"关键词提取失败: {str(e)}")

            # 备用方案：返回基于输入的问题的简单关键词
            words = question.split()
            en_keywords = words[:keywords_num] if len(words) >= keywords_num else words
            cn_keywords = words[:keywords_num] if len(words) >= keywords_num else words

            return {"keyword_en": en_keywords, "keyword_cn": cn_keywords, "searchquery_en": [question], "searchquery_cn": [question], "time_range": []}

    @classmethod
    def _parse_keyword_response(cls, response: str) -> Dict[str, List[str]]:
        """解析AI响应并提取关键词"""
        import re
        import json

        # 清理响应内容，移除代码块标记
        cleaned_response = response.strip()
        cleaned_response = re.sub(r"^```json\s*", "", cleaned_response)
        cleaned_response = re.sub(r"```\n\n    @classmethod\n    def _build_survey_prompt", "    @classmethod\n    def _build_survey_prompt", cleaned_response)
        cleaned_response = re.sub(r"```", "", cleaned_response)

        try:
            # 尝试解析JSON
            result = json.loads(cleaned_response)
            return {
                "keyword_en": result.get("keyword_en", []),
                "keyword_cn": result.get("keyword_cn", []),
                "searchquery_en": result.get("searchquery_en", []),
                "searchquery_cn": result.get("searchquery_cn", []),
                "time_range": result.get("time_range", []),
            }
        except json.JSONDecodeError:
            logger.warning(f"JSON解析失败，尝试提取关键词信息: {cleaned_response}")

            # 如果JSON解析失败，尝试用正则表达式提取信息
            keyword_en = re.findall(r'"keyword_en":\s*\[([^\]]+)\]', cleaned_response)
            keyword_cn = re.findall(r'"keyword_cn":\s*\[([^\]]+)\]', cleaned_response)
            searchquery_en = re.findall(r'"searchquery_en":\s*\[([^\]]+)\]', cleaned_response)
            searchquery_cn = re.findall(r'"searchquery_cn":\s*\[([^\]]+)\]', cleaned_response)
            time_range = re.findall(r'"time_range":\s*\[([^\]]+)\]', cleaned_response)

            # 简单处理提取的字符串
            def parse_list_str(list_str):
                if list_str:
                    items = [item.strip().strip("\"'") for item in list_str[0].split(",")]
                    return [item for item in items if item]
                return []

            return {
                "keyword_en": parse_list_str(keyword_en),
                "keyword_cn": parse_list_str(keyword_cn),
                "searchquery_en": parse_list_str(searchquery_en),
                "searchquery_cn": parse_list_str(searchquery_cn),
                "time_range": parse_list_str(time_range),
            }

    @classmethod
    def _build_survey_prompt(cls, papers: List[Dict[str, Any]]) -> str:
        """构建综述生成提示词"""
        logger.debug(f"收到{len(papers)}篇论文用于构建综述提示词")

        # 构建详细的论文信息
        papers_info = []
        titles = []

        for i, paper in enumerate(papers, 1):
            title = paper.get("title", "未知标题")
            abstract = paper.get("abstract", "")
            source = paper.get("source", "未知来源")
            similarity = paper.get("similarity", 0.0)
            # 获取论文全文或详细内容
            full_content = paper.get("full_content", "")

            titles.append(title)

            # 构建详细的论文信息，优先使用全文内容
            if full_content:
                paper_info = f"""
论文 {i}:
- 标题: {title}
- 来源: {source}
- 相似度: {similarity:.4f}
- content: {full_content[:2000]}...  # 限制长度以避免超限
"""
            else:
                # 如果没有全文content，则使用摘要
                paper_info = f"""
论文 {i}:
- 标题: {title}
- 来源: {source}
- 相似度: {similarity:.4f}
- 摘要: {abstract[:200]}...  # 限制摘要长度
"""
            papers_info.append(paper_info)

            logger.debug(f"第{i}篇论文: {title} | 来源: {source} | 相似度: {similarity:.4f}")

        papers_detail = "\n".join(papers_info)
        papers_titles_str = "、".join(titles)
        logger.debug(f"构建的论文标题字符串: '{papers_titles_str}'")

        return f"""
###目标###
我需要你作为文献综述专家，基于以下{len(papers)}篇学术论文的详细信息，生成一篇结构完整、content详实的中文文献综述。要求：
1. 深度分析每篇论文的核心观点、技术创新点和研究脉络
2. 按技术发展/应用场景/研究方法等维度组织content
3. 对每项具体技术或观点都要标注引用来源，使用 ##编号$ 格式
4. 保持学术严谨性，区分作者观点与综合评述

###指定论文详细信息###
{papers_detail}

###数据处理要求###
1. **深度内容提取**：
   - 提取每篇论文的核心技术方法、实验结果、创新点
   - 识别论文中的关键数据、性能指标、对比分析
   - 挖掘论文的技术局限性和未来展望

2. **跨文献对比分析**：
   - 对比不同论文的技术路线差异
   - 分析方法演进的时间脉络
   - 识别研究热点和技术趋势

###结构要求###
采用「总-分-总」结构，包含：
1. **研究背景**（融合{len(papers)}篇文献的研究动机和问题定义）
2. **核心技术进展分析**（按主题分类，深入对比各文献的方法创新）
   - 每个技术点必须引用具体论文content
   - 包含具体的技术细节、实验数据、性能对比
   - 分析不同方法的优缺点和适用场景
3. **技术挑战与局限**（结合各文献指出的问题和瓶颈）
4. **未来研究方向**（综合各文献的建议和展望）

###引用规范###
1. **引用密度**：每个技术点、观点或数据都必须关联≥1篇具体文献
2. **引用格式**：使用 ##编号$ 格式，编号从1开始按文献顺序
3. **content深度**：引用时要提及具体的技术方法、实验结果或理论观点
4. **全面覆盖**：确保每一篇指定论文都在综述中被充分分析

###质量标准###
1. **具体性**：避免泛泛而谈，要有具体的技术细节和数据支撑
2. **对比性**：横向对比不同论文的方法差异和性能优劣
3. **创新性**：突出每篇论文的独特贡献和技术突破
4. **批判性**：客观分析各方法的局限性和改进空间

###格式约束###
- 使用Markdown格式，标题层级清晰
- **引用格式**：必须使用 ##编号$$ 格式，编号范围1-{len(papers)}
- **严格禁止**：添加参考文献列表或References章节
- **严格禁止**：使用[1]、（1）、文献[1]等其他引用格式
- 每段content都要有具体的引用支撑

###引用编号对应关系###
请严格按照以下编号对应关系进行引用：
{chr(10).join([f"##{i + 1}$$ - {title}" for i, title in enumerate(titles)])}

###示例引用风格###
正确示例：
- "该研究提出的注意力机制在BLEU评分上达到了85.3% ##1$$，相比传统方法提升了12.7个百分点。"
- "##2$$ 的实验表明，多模态融合策略在视觉问答任务中的准确率为91.2%，但在推理时间上增加了约30%的开销。"
- "三种不同的技术路线显示出各自优势 ##1$$ ##2$$ ##3$$。"

错误示例：
- "多模态AI是重要研究方向。" （缺乏引用）
- "文献[1]指出..." （错误的引用格式）
- "根据文献1的研究..." （错误的引用格式）

请严格按照上述格式要求生成综述。
"""

    @classmethod
    def _build_survey_prompt_from_summaries(cls, paper_summaries: List[Dict[str, str]]) -> str:
        """基于简报构建综述生成提示词"""
        logger.debug(f"收到{len(paper_summaries)}份文献简报用于构建综述提示词")

        # 构建简报信息
        summaries_info = []
        titles = []

        for i, summary_info in enumerate(paper_summaries, 1):
            title = summary_info.get("title", "未知标题")
            summary = summary_info.get("summary", "")

            titles.append(title)

            # 构建简报信息
            summary_text = f"""
简报 {i}:
- 标题: {title}
- 简报: {summary[:2000]}...  # 限制长度以避免超限
"""
            summaries_info.append(summary_text)

            logger.debug(f"第{i}份简报: {title}")

        summaries_detail = "\n".join(summaries_info)
        summaries_titles_str = "、".join(titles)
        logger.debug(f"构建的简报标题字符串: '{summaries_titles_str}'")

        return f"""
###目标###
我需要你作为文献综述专家，基于以下{len(paper_summaries)}份文献简报，生成一篇结构完整、内容详实的中文文献综述。要求：
1. 深度分析每份简报对应论文的核心观点、技术创新点和研究脉络
2. 按技术发展/应用场景/研究方法等维度组织内容
3. 对每项具体技术或观点都要标注引用来源，使用 ##编号$ 格式
4. 保持学术严谨性，区分作者观点与综合评述

###指定文献简报内容###
{summaries_detail}

###数据处理要求###
1. **深度内容提取**：
   - 提取每篇论文的核心技术方法、实验结果、创新点
   - 识别论文中的关键数据、性能指标、对比分析
   - 挖掘论文的技术局限性和未来展望

2. **跨文献对比分析**：
   - 对比不同论文的技术路线差异
   - 分析方法演进的时间脉络
   - 识别研究热点和技术趋势

###结构要求###
采用「总-分-总」结构，包含：
1. **研究背景**（融合{len(paper_summaries)}篇文献的研究动机和问题定义）
2. **核心技术进展分析**（按主题分类，深入对比各文献的方法创新）
   - 每个技术点必须引用具体论文内容
   - 包含具体的技术细节、实验数据、性能对比
   - 分析不同方法的优缺点和适用场景
3. **技术挑战与局限**（结合各文献指出的问题和瓶颈）
4. **未来研究方向**（综合各文献的建议和展望）

###引用规范###
1. **引用密度**：每个技术点、观点或数据都必须关联≥1篇具体文献
2. **引用格式**：使用 ##编号$$ 格式，编号从1开始按文献顺序
3. **内容深度**：引用时要提及具体的技术方法、实验结果或理论观点
4. **全面覆盖**：确保每一份简报对应的论文都在综述中被充分分析

###质量标准###
1. **具体性**：避免泛泛而谈，要有具体的技术细节和数据支撑
2. **对比性**：横向对比不同论文的方法差异和性能优劣
3. **创新性**：突出每篇论文的独特贡献和技术突破
4. **批判性**：客观分析各方法的局限性和改进空间

###格式约束###
- 使用Markdown格式，标题层级清晰
- **引用格式**：必须使用 ##编号$$ 格式，编号范围1-{len(paper_summaries)}
- **严格禁止**：添加参考文献列表或References章节
- **严格禁止**：使用[1]、（1）、文献[1]等其他引用格式
- 每段内容都要有具体的引用支撑

###引用编号对应关系###
请严格按照以下编号对应关系进行引用：
{chr(10).join([f"##{i + 1}$$ - {title}" for i, title in enumerate(titles)])}

###示例引用风格###
正确示例：
- "该研究提出的注意力机制在BLEU评分上达到了85.3% ##1$$，相比传统方法提升了12.7个百分点。"
- "##2$$ 的实验表明，多模态融合策略在视觉问答任务中的准确率为91.2%，但在推理时间上增加了约30%的开销。"
- "三种不同的技术路线显示出各自优势 ##1$$ ##2$$ ##3$$。"

错误示例：
- "多模态AI是重要研究方向。" （缺乏引用）
- "文献[1]指出..." （错误的引用格式）
- "根据文献1的研究..." （错误的引用格式）

请严格按照上述格式要求生成综述。
"""

    @classmethod
    def _generate_word_document(cls, content: str, title: str, papers: list = []) -> bytes:
        """生成Word文档"""
        try:
            import re

            logger.info(f"开始生成Word文档: 标题={title}, 文献数量={len(papers)}, 内容长度={len(content)}")

            # 创建Word文档
            doc = Document()

            # 添加标题
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # 添加生成时间
            time_para = doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            time_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # 添加空行
            doc.add_paragraph("")

            # 收集所有引用编号
            cited_indices = set()
            citation_pattern = re.compile(r"##(\d+)\$\$")
            for match in citation_pattern.finditer(content):
                cited_indices.add(int(match.group(1)))

            logger.info(f"找到{len(cited_indices)}个引用: {sorted(cited_indices)}")

            # 处理并添加内容
            lines = content.split("\n")
            logger.info(f"处理{len(lines)}行内容...")

            for line in lines:
                line = line.strip()
                if not line:
                    continue

                # 检查是否为标题
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("- "):
                    # 项目符号列表
                    doc.add_paragraph(line[2:], style="List Bullet")
                else:
                    # 普通段落 - 需要处理引用标记
                    para = doc.add_paragraph()
                    cls._add_paragraph_with_citations(para, line)

            logger.info("内容添加完成,开始添加参考文献...")

            # 如果有引用,添加引用列表
            if papers and cited_indices:
                # 添加分隔线
                doc.add_paragraph("")
                doc.add_paragraph("─" * 50)
                doc.add_paragraph("")

                # 添加引用列表标题
                doc.add_heading("参考文献", level=2)

                # 按引用编号排序并添加文献信息
                sorted_indices = sorted(cited_indices)
                for idx in sorted_indices:
                    # papers是从0开始索引,引用编号从1开始
                    paper_idx = idx - 1
                    if 0 <= paper_idx < len(papers):
                        paper = papers[paper_idx]

                        # 只添加引用编号和标题
                        ref_para = doc.add_paragraph(style="List Number")
                        ref_para.add_run(f"[{idx}] ").bold = True
                        ref_para.add_run(paper.get("title", "未知标题"))

                logger.info(f"已添加{len(sorted_indices)}篇参考文献")
            else:
                if not papers:
                    logger.warning("没有文献列表,跳过参考文献章节")
                if not cited_indices:
                    logger.warning("内容中没有引用标记,跳过参考文献章节")

            # 保存到内存
            logger.info("开始保存文档到内存...")
            file_stream = BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            file_bytes = file_stream.read()
            logger.info(f"文档保存完成,大小: {len(file_bytes)} bytes")

            return file_bytes
        except Exception as e:
            logger.error(f"生成Word文档失败: {e}", exc_info=True)
            # 如果Word生成失败，返回纯文本
            text_content = f"{title}\n\n{content}\n\n生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            return text_content.encode("utf-8")

    @classmethod
    def _add_paragraph_with_citations(cls, para, text: str):
        """添加包含引用标记的段落"""
        import re

        # 匹配 ##数字$ 格式的引用
        citation_pattern = re.compile(r"##(\d+)\$\$")

        last_end = 0
        for match in citation_pattern.finditer(text):
            # 添加引用标记之前的文本
            if match.start() > last_end:
                para.add_run(text[last_end : match.start()])

            # 添加引用标记(上标格式)
            citation_run = para.add_run(f"[{match.group(1)}]")
            citation_run.font.superscript = True
            citation_run.font.color.rgb = RGBColor(0, 102, 204)  # 蓝色

            last_end = match.end()

        # 添加剩余文本
        if last_end < len(text):
            para.add_run(text[last_end:])
